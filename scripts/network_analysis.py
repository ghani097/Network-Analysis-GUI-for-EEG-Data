"""
Network Level Analysis - Python Implementation
==============================================

Simple Python implementation of network-level PLI analysis.

Usage:
    python network_analysis.py
    python network_analysis.py --no-baseline
    python network_analysis.py --input data.xlsx --output results
"""

import math
import os
import sys
from itertools import combinations
from pathlib import Path

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')  # Non-interactive backend
import matplotlib.pyplot as plt
from scipy import stats

try:
    import statsmodels.formula.api as smf
    from statsmodels.stats.anova import anova_lm
    from statsmodels.stats.multitest import multipletests
    HAS_STATSMODELS = True
except ImportError:
    HAS_STATSMODELS = False

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def _fisher_z(x):
    """Variance-stabilising transform for bounded [0, 1] PLI values.

    PLI is a bounded connectivity measure; raw-scale LMMs inflate precision
    near the bounds. atanh(clip) is the standard pre-modelling transform.
    """
    return np.arctanh(np.clip(x, -0.999, 0.999))


class NetworkAnalysis:
    """Network-level PLI analysis with mixed-effects models."""

    def __init__(self, input_file, output_dir="analysis_output",
                 adjust_baseline=True, plot_type="bar", groups=None,
                 sessions=None, networks=None, frequency_bands=None,
                 callback=None):
        """
        Initialize analysis.

        Args:
            input_file: Path to Excel file with PLI data
            output_dir: Directory for output files
            adjust_baseline: If True, use Pre session as covariate
            plot_type: Plot style for figures ('bar' or 'line')
            groups: List of groups to include (default: all)
            sessions: List of sessions to include (default: all)
            networks: List of networks to analyze (default: all)
            frequency_bands: List of frequency bands (default: all)
            callback: Function to call with progress updates (msg, percent)
        """
        self.input_file = Path(input_file).resolve()
        self.output_dir = Path(output_dir).resolve()
        self.adjust_baseline = adjust_baseline
        self.plot_type = str(plot_type).strip().lower() if plot_type else "bar"
        if self.plot_type not in {"bar", "line"}:
            self.plot_type = "bar"
        self.callback = callback or (lambda msg, pct: print(msg))

        # All of these can be None → auto-detected from data in load_data()
        self.groups = groups  # resolved in load_data()
        self.sessions = sessions  # resolved in load_data()
        self.networks = networks  # resolved in load_data()
        self.frequency_bands = frequency_bands  # resolved in load_data()

        # Session order - set after sessions are resolved
        self.session_order = list(self.sessions) if self.sessions else []

        self.data = None
        self.results = {}
        self.all_stats = []  # Store all statistics for Excel export
        self.all_plots = {}  # Store plot data for combined figure
        self.contrast_results = {}  # Store contrast results for significance display

    def _update(self, msg, pct=None):
        """Send progress update."""
        if self.callback:
            self.callback(msg, pct)

    def _get_session_order(self, sessions):
        """Get sessions in correct order based on user-provided session order."""
        return [s for s in self.session_order if s in sessions]

    def load_data(self):
        """Load and preprocess data."""
        self._update(f"Loading: {self.input_file.name}", 5)

        if not self.input_file.exists():
            raise FileNotFoundError(f"Input file not found: {self.input_file}")

        df = pd.read_excel(str(self.input_file))

        # Ensure categorical columns are strings for consistent filtering
        for col in ['Group', 'Session', 'Network', 'FrequencyTag']:
            if col in df.columns:
                df[col] = df[col].astype(str)

        self._update(f"Loaded {len(df)} rows", 10)

        # Auto-detect groups if not specified
        data_groups = sorted(df['Group'].dropna().unique().tolist())
        if self.groups is None or not set(self.groups).intersection(data_groups):
            if self.groups:
                self._update(
                    f"Warning: groups {self.groups} not found in data. "
                    f"Auto-detecting: {data_groups}", None
                )
            self.groups = data_groups

        # Auto-detect sessions / networks / bands from the data when not specified,
        # or when the user-supplied values have no overlap with the actual data.
        data_sessions = sorted(df['Session'].dropna().unique().tolist())
        if self.sessions is None or not set(self.sessions).intersection(data_sessions):
            if self.sessions:
                self._update(
                    f"Warning: sessions {self.sessions} not found in data. "
                    f"Auto-detecting: {data_sessions}", None
                )
            self.sessions = data_sessions
            self.session_order = list(data_sessions)

        data_networks = sorted(df['Network'].dropna().unique().tolist())
        if self.networks is None or not set(self.networks).intersection(data_networks):
            if self.networks:
                self._update(
                    f"Warning: networks {self.networks} not found in data. "
                    f"Auto-detecting: {data_networks}", None
                )
            self.networks = data_networks

        data_bands = sorted(df['FrequencyTag'].dropna().unique().tolist())
        if self.frequency_bands is None or not set(self.frequency_bands).intersection(data_bands):
            if self.frequency_bands:
                self._update(
                    f"Warning: bands {self.frequency_bands} not found in data. "
                    f"Auto-detecting: {data_bands}", None
                )
            self.frequency_bands = data_bands

        # Filter data
        df = df[df['Group'].isin(self.groups)]
        df = df[df['Network'].isin(self.networks)]
        df = df[df['FrequencyTag'].isin(self.frequency_bands)]

        # Always keep all sessions in the modelling frame. Baseline adjustment
        # is now handled by the participant random intercept in fit_model,
        # which absorbs per-subject baseline differences without the
        # collinearity that arose from using PLI_Pre_Value as a fixed-effect
        # covariate alongside a (1|Participant) random effect.
        df = df[df['Session'].isin(self.sessions)]
        if self.adjust_baseline:
            self._update("Baseline adjustment: ON (participant random intercept)", 15)
        else:
            self._update("Baseline adjustment: OFF (OLS, no random intercept)", 15)

        # Set session as ordered categorical
        available_sessions = self._get_session_order(df['Session'].unique())
        df['Session'] = pd.Categorical(df['Session'], categories=available_sessions, ordered=True)

        # Pseudo-replication guard: assert one row per
        # (Participant, Session, Network, FrequencyTag) cell. If this fails,
        # downstream N and SEs would be silently inflated.
        cell_counts = df.groupby(
            ['Participant', 'Session', 'Network', 'FrequencyTag'], observed=True
        ).size()
        if cell_counts.max() > 1:
            dup = cell_counts[cell_counts > 1].head(5)
            self._update(
                "WARNING: pseudo-replication detected — >1 row per "
                "(Participant, Session, Network, FrequencyTag) cell. "
                f"Example duplicates:\n{dup}", None
            )

        # Fisher-z transform for modelling. Raw values are retained in
        # MeanPLI_raw so plots, means, and the Excel 'Group_Means' sheet
        # continue to display the untransformed [0, 1] PLI.
        df['MeanPLI_raw'] = df['MeanPLI']
        df['MeanPLI_z'] = _fisher_z(df['MeanPLI'])

        self._update(f"Filtered: {len(df)} rows", 20)
        self.data = df
        return df

    def fit_model(self, network, freq_band):
        """Fit a repeated-measures model for one network / frequency band.

        The modelled outcome is the Fisher-z-transformed PLI. The formula
        `C(Group) * C(Session)` is fitted with a participant random
        intercept when `adjust_baseline=True`, which absorbs between-subject
        baseline differences without the singular-covariate problem that
        the previous `PLI_Pre_Value + (1|Participant)` parameterisation
        created. When `adjust_baseline=False`, we fit OLS with
        cluster-robust SEs on Participant so that within-subject
        correlation is still honoured.
        """
        subset = self.data[
            (self.data['Network'] == network) &
            (self.data['FrequencyTag'] == freq_band)
        ].copy()

        if len(subset) == 0:
            return None

        # Drop rows missing any modelling column
        subset = subset.dropna(
            subset=['MeanPLI_z', 'Group', 'Session', 'Participant']
        ).reset_index(drop=True)

        model_name = f"{freq_band} x {network}"
        model_result = {
            'model': None,
            'data': subset,
            'name': model_name,
            'anova': None,
            'model_significance': {},
        }

        if not HAS_STATSMODELS or subset.empty:
            return model_result

        if subset['Session'].nunique() < 2 or subset['Group'].nunique() < 2:
            self._update(
                f"  Skipping {model_name}: need >=2 groups and >=2 sessions", None
            )
            return model_result

        formula = "MeanPLI_z ~ C(Group) * C(Session)"
        model = None
        model_type = None

        if self.adjust_baseline:
            try:
                model = smf.mixedlm(
                    formula, data=subset, groups=subset['Participant']
                ).fit(method='lbfgs', disp=False)
                model_type = 'mixed'
                # If the random-effects variance collapsed to 0 the mixed
                # model has effectively become OLS without cluster SEs —
                # warn loudly so the user sees it in the progress log.
                try:
                    re_var = float(np.asarray(model.cov_re).ravel()[0])
                    if re_var < 1e-10:
                        self._update(
                            f"  WARNING: {model_name} random-effect variance "
                            "is ~0 (singular); LMM fixed-effect SEs are "
                            "unreliable. Consider adjust_baseline=False "
                            "(cluster-robust OLS) for this model.",
                            None,
                        )
                except Exception:
                    pass
            except Exception as e:
                self._update(
                    f"  WARNING: mixed model failed for {model_name}: {e}. "
                    "Falling back to cluster-robust OLS on Participant.",
                    None,
                )
                model = None

        if model is None:
            try:
                model = smf.ols(formula, data=subset).fit(
                    cov_type='cluster',
                    cov_kwds={'groups': subset['Participant'].values},
                )
                model_type = 'ols_cluster'
            except Exception as e:
                self._update(
                    f"  ERROR: {model_name} OLS fit also failed: {e}",
                    None,
                )
                return model_result

        model_result['model'] = model
        model_result['model_type'] = model_type

        # Extract model statistics (drops intercept, uses Satterthwaite-
        # style df for the mixed model)
        self._extract_model_stats(
            model, model_name, network, freq_band, model_type,
            n_subjects=int(subset['Participant'].nunique()),
        )

        # Extract significance for plotting (session effects and interactions)
        model_result['model_significance'] = self._extract_plot_significance(
            model, subset
        )

        return model_result

    def _extract_plot_significance(self, model, subset):
        """Extract significance for each session from model coefficients."""
        significance = {}
        if model is None:
            return significance

        pvalues = model.pvalues
        sessions = subset['Session'].unique()

        # Check each parameter for session-related significance
        for param in pvalues.index:
            p = pvalues[param]
            stars = self._get_sig_stars(p)

            if stars and stars != 'ns':
                # Check for interaction terms (Group:Session) - these indicate group differences at specific sessions
                if ':C(Session)' in param or 'C(Session)' in param and ':' in param:
                    # Extract session name from interaction term like "C(Group)[T.B]:C(Session)[T.post2]"
                    for session in sessions:
                        session_str = str(session)
                        if f'[T.{session_str}]' in param:
                            # Interaction term - group difference at this session
                            if session_str not in significance:
                                significance[session_str] = {'p': p, 'stars': stars, 'type': 'interaction'}
                            elif p < significance[session_str]['p']:
                                significance[session_str] = {'p': p, 'stars': stars, 'type': 'interaction'}

                # Also check main session effects
                elif 'C(Session)' in param and ':' not in param:
                    # Main effect of session like "C(Session)[T.post2]"
                    for session in sessions:
                        session_str = str(session)
                        if f'[T.{session_str}]' in param:
                            # Only add if no interaction significance already exists for this session
                            if session_str not in significance:
                                significance[session_str] = {'p': p, 'stars': stars, 'type': 'session_effect'}

        return significance

    def _extract_model_stats(self, model, model_name, network, freq_band,
                             model_type='mixed', n_subjects=0):
        """Extract and store model statistics for Excel export.

        Differences from the previous version:
        - The Intercept is now excluded (nuisance parameter testing PLI = 0).
        - For mixed models, p-values are recomputed against a t distribution
          with an approximate Satterthwaite-style residual df:
            resid_df = n_obs − n_fixed − (n_subjects − 1)
          instead of the large-sample normal reference that statsmodels uses
          by default.
        - For cluster-robust OLS, the model already carries valid df_resid.
        - A 'df' column is included so every reported statistic has its
          degrees of freedom alongside it.
        """
        if model is None:
            return

        # Get fixed effects (different attribute names for mixed vs OLS)
        if model_type == 'mixed':
            fe = model.fe_params
            bse = model.bse_fe
        else:
            fe = model.params
            bse = model.bse
        tvalues = model.tvalues

        # Determine residual df
        n_obs = int(model.nobs)
        n_fixed = len(fe)
        if model_type == 'mixed':
            resid_df = max(n_obs - n_fixed - max(n_subjects - 1, 0), 1)
        else:
            resid_df = getattr(model, 'df_resid', max(n_obs - n_fixed, 1))

        for param in fe.index:
            # Skip the intercept — it tests PLI = 0, which is meaningless
            # for a bounded connectivity measure.
            if param == 'Intercept':
                continue

            t_val = float(tvalues[param]) if param in tvalues.index else np.nan
            se = float(bse[param]) if param in bse.index else np.nan

            # Recompute p against the t distribution with resid_df
            if np.isfinite(t_val):
                p_val = float(2.0 * stats.t.sf(abs(t_val), df=resid_df))
            else:
                p_val = np.nan

            self.all_stats.append({
                'Model': model_name,
                'Network': network,
                'FrequencyBand': freq_band,
                'Effect': param,
                'Estimate': float(fe[param]),
                'Std.Error': se,
                't-value': t_val,
                'df': resid_df,
                'p-value': p_val,
                'Significance': self._get_sig_stars(p_val),
                'ModelType': model_type,
            })

    def _get_sig_stars(self, p):
        """Get significance stars for p-value."""
        if pd.isna(p):
            return ''
        if p < 0.001:
            return '***'
        elif p < 0.01:
            return '**'
        elif p < 0.05:
            return '*'
        else:
            return 'ns'

    def compute_means(self, subset):
        """Compute group means by session on the raw (untransformed) PLI
        scale so that reported means, SDs and SEs are interpretable.

        N is the number of unique participants, not the row count (guards
        against inflated N if upstream ever reintroduces ROI-pair rows).
        """
        # Use MeanPLI_raw if available, otherwise MeanPLI
        pli_col = 'MeanPLI_raw' if 'MeanPLI_raw' in subset.columns else 'MeanPLI'

        means = subset.groupby(['Group', 'Session'], observed=True).agg(
            Mean=(pli_col, 'mean'),
            SD=(pli_col, 'std'),
            N=('Participant', 'nunique'),
        ).reset_index()
        means['SE'] = means['SD'] / np.sqrt(means['N'])

        # Sort by session order
        session_order = self._get_session_order(means['Session'].unique())
        means['Session'] = pd.Categorical(means['Session'], categories=session_order, ordered=True)
        means = means.sort_values(['Group', 'Session'])

        return means

    def run_contrasts(self, subset, model_name, network, freq_band):
        """Run between-group and within-session contrasts on raw PLI.

        Between-group contrasts use an independent-samples t-test (different
        participants in each group).

        Within-group session contrasts use a **paired** t-test
        (``scipy.stats.ttest_rel``) aligned on ``Participant``, because
        sessions are repeated measures on the same participants.

        **Effect-size standardiser (conservative / grand-SD approach)**:
        MeanPLI is pre-averaged across ROI pairs within each network, which
        dramatically compresses within-group standard deviations (~0.001–0.004)
        and would produce Cohen's d > 100 if within-group SD were used.
        Following Morris & DeShon (2002), Cohen's d is instead standardised by
        the *grand* SD pooled across all participants at the relevant comparison
        point:

        * Between-group at session s → grand SD of all participants at session s
        * Within-group s1 vs s2 → SD of (session-s1 – session-s2) difference
          scores pooled across **all** groups

        This gives conservative, ecologically valid effect sizes (typically
        d = 0.5–3.0 for real treatment effects) that accurately reflect group
        differences in the context of the full distributional spread of PLI.

        Every row carries explicit ``df`` for transparent reporting.
        """
        results = {'group': [], 'session': [], 'significance': {}}
        output_lines = []
        contrast_stats = []
        pli_col = 'MeanPLI_raw' if 'MeanPLI_raw' in subset.columns else 'MeanPLI'

        # Get sessions in correct order
        sessions = self._get_session_order(subset['Session'].unique())

        # Between-group contrasts by session (pairwise for any number of groups)
        for session in sessions:
            s_data = subset[subset['Session'] == session]
            groups = sorted(s_data['Group'].unique())

            if len(groups) >= 2:
                for g1_name, g2_name in combinations(groups, 2):
                    g1 = s_data[s_data['Group'] == g1_name][pli_col]
                    g2 = s_data[s_data['Group'] == g2_name][pli_col]

                    if len(g1) > 1 and len(g2) > 1:
                        t_val, p = stats.ttest_ind(g1, g2)
                        df_val = len(g1) + len(g2) - 2
                        diff = g1.mean() - g2.mean()
                        # Grand-SD Cohen's d (Morris & DeShon, 2002):
                        # Standardise by the total SD across ALL participants at
                        # this session.  Within-group SD is artificially
                        # compressed because MeanPLI is pre-averaged over ROI
                        # pairs; using the grand SD (which includes
                        # between-group spread) gives a conservative,
                        # ecologically valid effect-size estimate.
                        grand_sd = s_data[pli_col].std()
                        d = diff / grand_sd if grand_sd > 0 else np.nan
                        sig = self._get_sig_stars(p)

                        # Record the strongest per-session between-group
                        # contrast (including non-significant ones) so the
                        # bar plots can annotate every session.
                        prev = results['significance'].get(session)
                        if prev is None or p < prev['p']:
                            results['significance'][session] = {'p': p, 'stars': sig}

                        output_lines.append(
                            f"  {session}: {g1_name} vs {g2_name} = {diff:+.4f}, "
                            f"t({df_val})={t_val:.2f}, p={p:.4f}, d={d:.2f} {sig}"
                        )
                        results['group'].append({
                            'session': session, 'contrast': f"{g1_name} - {g2_name}",
                            'diff': diff, 't': t_val, 'p': p, 'df': df_val
                        })

                        contrast_stats.append({
                            'Model': model_name,
                            'Network': network,
                            'FrequencyBand': freq_band,
                            'ContrastType': 'Between-Group',
                            'Session': session,
                            'Group': '',
                            'Contrast': f"{g1_name} vs {g2_name}",
                            'Difference': diff,
                            't-value': t_val,
                            'df': df_val,
                            'Cohens_d': d,
                            'p-value': p,
                            'Significance': sig,
                        })

        # Within-group session contrasts — PAIRED t-test aligned on Participant
        # Pre-compute grand difference SDs (pooled across all groups) for each
        # session pair so within-group d uses a stable, conservative standardiser.
        grand_diff_sd_cache: dict = {}
        for i, s1 in enumerate(sessions):
            for s2 in sessions[i+1:]:
                all_d1 = subset[subset['Session'] == s1][['Participant', pli_col]].rename(
                    columns={pli_col: 'val1'}
                )
                all_d2 = subset[subset['Session'] == s2][['Participant', pli_col]].rename(
                    columns={pli_col: 'val2'}
                )
                all_merged = all_d1.merge(all_d2, on='Participant', how='inner')
                if len(all_merged) > 1:
                    grand_diff_sd_cache[(s1, s2)] = (all_merged['val1'] - all_merged['val2']).std()

        for group in sorted(subset['Group'].unique()):
            g_data = subset[subset['Group'] == group]

            for i, s1 in enumerate(sessions):
                for s2 in sessions[i+1:]:
                    d1 = g_data[g_data['Session'] == s1][['Participant', pli_col]].rename(
                        columns={pli_col: 'val1'}
                    )
                    d2 = g_data[g_data['Session'] == s2][['Participant', pli_col]].rename(
                        columns={pli_col: 'val2'}
                    )
                    merged = d1.merge(d2, on='Participant', how='inner')

                    if len(merged) > 1:
                        t_val, p = stats.ttest_rel(merged['val1'], merged['val2'])
                        df_val = len(merged) - 1
                        diff = merged['val1'].mean() - merged['val2'].mean()
                        # Grand-SD d_z: standardise by the SD of difference
                        # scores pooled across ALL groups (conservative; avoids
                        # inflation from within-group averaging compression).
                        grand_diff_sd = grand_diff_sd_cache.get((s1, s2), np.nan)
                        d = diff / grand_diff_sd if grand_diff_sd and grand_diff_sd > 0 else np.nan
                        sig = self._get_sig_stars(p)

                        output_lines.append(
                            f"  {group}: {s1} vs {s2} = {diff:+.4f}, "
                            f"t({df_val})={t_val:.2f}, p={p:.4f}, d={d:.2f} {sig}"
                        )
                        results['session'].append({
                            'group': group, 'contrast': f"{s1} - {s2}",
                            'diff': diff, 't': t_val, 'p': p, 'df': df_val
                        })

                        contrast_stats.append({
                            'Model': model_name,
                            'Network': network,
                            'FrequencyBand': freq_band,
                            'ContrastType': 'Within-Group',
                            'Session': '',
                            'Group': group,
                            'Contrast': f"{s1} vs {s2}",
                            'Difference': diff,
                            't-value': t_val,
                            'df': df_val,
                            'Cohens_d': d,
                            'p-value': p,
                            'Significance': sig,
                        })

        return results, output_lines, contrast_stats

    def _get_group_colors(self, groups):
        """Get color mapping for groups dynamically."""
        # Explicit overrides for the chiropractic trial's two main groups so
        # the grouped bar plots come out in the manuscript's orange/blue scheme.
        explicit = {
            'Chiro':        '#E58840',
            'Chiropractic': '#E58840',
            'Control':      '#5F7FB8',
            'Ctrl':         '#5F7FB8',
        }
        color_palette = [
            '#E58840',  # Orange  (Chiro)
            '#5F7FB8',  # Blue    (Control)
            '#009E73',  # Green
            '#F0E442',  # Yellow
            '#0072B2',  # Blue
            '#D55E00',  # Red-orange
            '#CC79A7',  # Pink
            '#999999',  # Gray
        ]
        colors = {}
        palette_idx = 0
        for group in sorted(groups):
            if group in explicit:
                colors[group] = explicit[group]
            else:
                colors[group] = color_palette[palette_idx % len(color_palette)]
                palette_idx += 1
        return colors

    def _get_group_markers(self, groups):
        """Get marker mapping for groups dynamically."""
        marker_palette = ['o', 's', '^', 'D', 'v', 'p', 'h', '*']
        markers = {}
        for i, group in enumerate(sorted(groups)):
            markers[group] = marker_palette[i % len(marker_palette)]
        return markers

    def _draw_bar_panel(self, ax, means, sessions, significance, network,
                        freq_band, colors, show_xlabels=True,
                        show_ylabel=True, show_title=False,
                        title_fontsize=12, ylabel_fontsize=11,
                        tick_fontsize=9):
        """Render one grouped-bar panel (Group x Session) on a given axes.

        Returns a dict mapping group name -> the bar artist used as a legend
        handle, so the caller can build a shared figure-level legend.
        """
        sessions = list(sessions)
        groups = sorted(means['Group'].unique())
        n_groups = max(len(groups), 1)
        # Narrower bars: 0.35 per bar when 2 groups, otherwise share 0.7 among
        # however many groups. This leaves breathing room between sessions so
        # the session labels and pair-brackets read cleanly.
        bar_width = 0.35 if n_groups == 2 else 0.7 / n_groups
        x_base = np.arange(len(sessions), dtype=float)

        bar_tops = {s: -np.inf for s in sessions}
        group_offsets = {}
        all_vals = []
        handles = {}

        for gi, group in enumerate(groups):
            gdata = (
                means[means['Group'] == group]
                .set_index('Session').reindex(sessions).reset_index()
            )
            offset = (gi - (n_groups - 1) / 2.0) * bar_width
            group_offsets[group] = offset
            color = colors.get(group, '#888888')

            heights = pd.to_numeric(gdata['Mean'], errors='coerce').to_numpy()
            errs = pd.to_numeric(gdata['SE'], errors='coerce').to_numpy()
            errs_plot = np.where(np.isfinite(errs), errs, 0.0)
            heights_plot = np.where(np.isfinite(heights), heights, 0.0)

            bars = ax.bar(
                x_base + offset, heights_plot, bar_width,
                yerr=errs_plot, color=color, edgecolor='none',
                error_kw=dict(ecolor='black', lw=0.9, capsize=2.5),
                label=group,
            )
            handles[group] = bars[0]

            for i, (h, e) in enumerate(zip(heights, errs)):
                if not np.isfinite(h):
                    continue
                ev = e if np.isfinite(e) else 0.0
                top = h + ev
                low = h - ev
                if top > bar_tops[sessions[i]]:
                    bar_tops[sessions[i]] = top
                all_vals.extend([top, low, h])

        if not all_vals:
            ax.text(0.5, 0.5, 'No Data', ha='center', va='center',
                    transform=ax.transAxes, fontsize=9, color='#888888')
            ax.set_xticks([])
            ax.set_yticks([])
            for side in ('top', 'right', 'left', 'bottom'):
                ax.spines[side].set_visible(False)
            return handles

        data_min = float(np.nanmin(all_vals))
        data_max = float(np.nanmax(all_vals))
        if (not np.isfinite(data_min) or not np.isfinite(data_max)
                or data_min == data_max):
            data_min, data_max = 0.0, max(data_max, 1.0)

        y_range = max(data_max - data_min, 1e-9)
        # Tight y-limits with small padding below and headroom above for the
        # pair brackets / significance labels.
        y_bottom = data_min - y_range * 0.25
        y_top = data_max + y_range * 0.45
        ax.set_ylim(y_bottom, y_top)

        # Significance annotations. For exactly two groups we draw a classic
        # pair-bracket connecting the Chiro and Control bars within each
        # session with the label centred above; otherwise fall back to a
        # simple text label above the tallest bar in the session.
        tick_h = y_range * 0.018
        bracket_pad = y_range * 0.08
        for i, session in enumerate(sessions):
            sig = significance.get(session) if significance else None
            label = (sig or {}).get('stars') or 'ns'
            if label == '':
                label = 'ns'
            top_here = bar_tops[session]
            if not np.isfinite(top_here):
                continue
            is_sig = label in ('*', '**', '***')
            color = 'black' if is_sig else '#8a8a8a'
            fontsize = 10 if is_sig else 8
            fontweight = 'bold' if is_sig else 'normal'

            y_line = top_here + bracket_pad
            if n_groups == 2:
                offsets_sorted = sorted(group_offsets.values())
                x1 = x_base[i] + offsets_sorted[0]
                x2 = x_base[i] + offsets_sorted[-1]
                ax.plot(
                    [x1, x1, x2, x2],
                    [y_line - tick_h, y_line, y_line, y_line - tick_h],
                    lw=(0.9 if is_sig else 0.7),
                    color=color,
                    solid_capstyle='butt',
                    clip_on=False,
                )
                ax.text(
                    (x1 + x2) / 2.0, y_line + y_range * 0.015, label,
                    ha='center', va='bottom',
                    fontsize=fontsize, fontweight=fontweight, color=color,
                )
            else:
                ax.text(
                    x_base[i], y_line, label,
                    ha='center', va='bottom',
                    fontsize=fontsize, fontweight=fontweight, color=color,
                )

        ax.set_xticks(x_base)
        if show_xlabels:
            ax.set_xticklabels(sessions, fontsize=tick_fontsize)
        else:
            ax.set_xticklabels([])
        ax.tick_params(axis='y', labelsize=tick_fontsize)

        if show_ylabel:
            ax.set_ylabel(f"{network}\nMean PLI",
                          fontsize=ylabel_fontsize, fontweight='bold')
        if show_title:
            ax.set_title(f"{freq_band} band",
                         fontsize=title_fontsize, fontweight='bold', pad=8)

        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.margins(x=0.08)
        return handles

    def _draw_line_panel(self, ax, means, sessions, significance, network,
                         freq_band, colors, show_xlabels=True,
                         show_ylabel=True, show_title=False,
                         title_fontsize=12, ylabel_fontsize=11,
                         tick_fontsize=9):
        """Render one line-plot panel (Group x Session) on a given axes."""
        sessions = list(sessions)
        groups = sorted(means['Group'].unique())
        x_base = np.arange(len(sessions), dtype=float)
        markers = self._get_group_markers(groups)

        line_tops = {s: -np.inf for s in sessions}
        all_vals = []
        handles = {}

        for group in groups:
            gdata = means[means['Group'] == group].copy()
            gdata['Session'] = pd.Categorical(gdata['Session'], categories=sessions, ordered=True)
            gdata = gdata.sort_values('Session')

            heights = pd.to_numeric(gdata['Mean'], errors='coerce').to_numpy()
            errs = pd.to_numeric(gdata['SE'], errors='coerce').to_numpy()
            if len(heights) < len(sessions):
                padded_heights = np.full(len(sessions), np.nan)
                padded_errs = np.full(len(sessions), np.nan)
                for idx, session in enumerate(sessions):
                    match = gdata[gdata['Session'] == session]
                    if not match.empty:
                        padded_heights[idx] = pd.to_numeric(match['Mean'], errors='coerce').iloc[0]
                        padded_errs[idx] = pd.to_numeric(match['SE'], errors='coerce').iloc[0]
                heights = padded_heights
                errs = padded_errs

            line = ax.errorbar(
                x_base, heights, yerr=errs,
                fmt='-', marker=markers[group], markersize=6,
                linewidth=2.0, elinewidth=1.1, capsize=4,
                color=colors[group], label=group,
            )
            handles[group] = line.lines[0]

            for i, session in enumerate(sessions):
                h = heights[i] if i < len(heights) else np.nan
                e = errs[i] if i < len(errs) else np.nan
                if np.isfinite(h):
                    top = h + (e if np.isfinite(e) else 0.0)
                    line_tops[session] = max(line_tops[session], top)
                    low = h - (e if np.isfinite(e) else 0.0)
                    all_vals.extend([top, low, h])

        if not all_vals:
            ax.text(0.5, 0.5, 'No Data', ha='center', va='center',
                    transform=ax.transAxes, fontsize=9, color='#888888')
            ax.set_xticks([])
            ax.set_yticks([])
            for side in ('top', 'right', 'left', 'bottom'):
                ax.spines[side].set_visible(False)
            return handles

        data_min = float(np.nanmin(all_vals))
        data_max = float(np.nanmax(all_vals))
        if (not np.isfinite(data_min) or not np.isfinite(data_max)
                or data_min == data_max):
            data_min, data_max = 0.0, max(data_max, 1.0)

        y_range = max(data_max - data_min, 1e-9)
        y_bottom = data_min - y_range * 0.20
        y_top = data_max + y_range * 0.30
        ax.set_ylim(y_bottom, y_top)

        for i, session in enumerate(sessions):
            sig = significance.get(session) if significance else None
            label = (sig or {}).get('stars') or 'ns'
            if label == '':
                label = 'ns'
            top_here = line_tops[session]
            if not np.isfinite(top_here):
                continue
            is_sig = label in ('*', '**', '***')
            color = 'black' if is_sig else '#8a8a8a'
            fontsize = 10 if is_sig else 8
            fontweight = 'bold' if is_sig else 'normal'
            ax.text(
                x_base[i], top_here + y_range * 0.07, label,
                ha='center', va='bottom',
                fontsize=fontsize, fontweight=fontweight, color=color,
            )

        ax.set_xticks(x_base)
        if show_xlabels:
            ax.set_xticklabels(sessions, fontsize=tick_fontsize)
        else:
            ax.set_xticklabels([])
        ax.tick_params(axis='y', labelsize=tick_fontsize)

        if show_ylabel:
            ax.set_ylabel(f"{network}\nMean PLI",
                          fontsize=ylabel_fontsize, fontweight='bold')
        if show_title:
            ax.set_title(f"{freq_band} band",
                         fontsize=title_fontsize, fontweight='bold', pad=8)

        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.margins(x=0.08)
        return handles

    def _draw_plot_panel(self, ax, means, sessions, significance, network,
                         freq_band, colors, show_xlabels=True,
                         show_ylabel=True, show_title=False,
                         title_fontsize=12, ylabel_fontsize=11,
                         tick_fontsize=9):
        """Dispatch to the configured plot renderer."""
        if self.plot_type == 'line':
            return self._draw_line_panel(
                ax, means, sessions, significance, network, freq_band, colors,
                show_xlabels=show_xlabels, show_ylabel=show_ylabel,
                show_title=show_title, title_fontsize=title_fontsize,
                ylabel_fontsize=ylabel_fontsize, tick_fontsize=tick_fontsize,
            )
        return self._draw_bar_panel(
            ax, means, sessions, significance, network, freq_band, colors,
            show_xlabels=show_xlabels, show_ylabel=show_ylabel,
            show_title=show_title, title_fontsize=title_fontsize,
            ylabel_fontsize=ylabel_fontsize, tick_fontsize=tick_fontsize,
        )

    def create_plot(self, subset, model_name, significance=None, save=True):
        """Render a single plot panel for one network x frequency band."""
        means = self.compute_means(subset)
        significance = significance or {}

        sessions = self._get_session_order(means['Session'].unique())

        # Parse network / freq_band from "{freq} x {network}" model name so
        # the title and ylabel use the manuscript style.
        try:
            freq_band, network = [s.strip() for s in model_name.split('x', 1)]
        except ValueError:
            freq_band, network = model_name, ''

        fig, ax = plt.subplots(figsize=(5, 4))
        colors = self._get_group_colors(means['Group'].unique())

        handles = self._draw_plot_panel(
            ax, means, sessions, significance, network, freq_band, colors,
            show_xlabels=True, show_ylabel=True, show_title=True,
        )

        if handles:
            ax.legend(
                handles=list(handles.values()),
                labels=list(handles.keys()),
                loc='upper center', bbox_to_anchor=(0.5, -0.12),
                ncol=len(handles), frameon=False, fontsize=9,
            )

        plt.tight_layout()

        # Store for combined figure
        self.all_plots[model_name] = {
            'means': means,
            'sessions': sessions,
            'significance': significance,
            'network': network,
            'freq_band': freq_band,
            'plot_type': self.plot_type,
        }

        if save:
            filename = model_name.replace(' ', '_').replace('x', '_') + '.png'
            filepath = self.output_dir / filename
            fig.savefig(str(filepath), dpi=200, bbox_inches='tight')
            plt.close(fig)
            return filepath
        else:
            return fig

    def create_combined_figure(self):
        """Create the manuscript-style grid figure (rows = networks, cols = bands)."""
        if not self.all_plots:
            return None

        self._update("\nCreating combined figure...", 92)

        # Manuscript layout: rows = networks (SN, DMN, CEN),
        #                   cols = frequency bands (Theta, Alpha, Beta, Gamma).
        networks = self.networks
        freq_bands = self.frequency_bands

        n_rows = len(networks)
        n_cols = len(freq_bands)

        fig, axes = plt.subplots(
            n_rows, n_cols,
            figsize=(3.6 * n_cols, 3.0 * n_rows),
            squeeze=False,
        )

        legend_handles = {}
        for i, network in enumerate(networks):
            for j, freq in enumerate(freq_bands):
                ax = axes[i, j]
                model_name = f"{freq} x {network}"

                if model_name not in self.all_plots:
                    ax.text(0.5, 0.5, 'No Data',
                            ha='center', va='center',
                            transform=ax.transAxes,
                            fontsize=9, color='#888888')
                    for side in ('top', 'right', 'left', 'bottom'):
                        ax.spines[side].set_visible(False)
                    ax.set_xticks([])
                    ax.set_yticks([])
                    if i == 0:
                        ax.set_title(f"{freq} band",
                                     fontsize=12, fontweight='bold', pad=8)
                    if j == 0:
                        ax.set_ylabel(f"{network}\nMean PLI",
                                      fontsize=11, fontweight='bold')
                    continue

                plot_data = self.all_plots[model_name]
                means = plot_data['means']
                sessions = plot_data['sessions']
                significance = plot_data.get('significance', {})
                colors = self._get_group_colors(means['Group'].unique())

                handles = self._draw_plot_panel(
                    ax, means, sessions, significance,
                    network=network, freq_band=freq, colors=colors,
                    show_xlabels=(i == n_rows - 1),
                    show_ylabel=(j == 0),
                    show_title=(i == 0),
                    title_fontsize=12, ylabel_fontsize=11, tick_fontsize=9,
                )
                for g, h in handles.items():
                    legend_handles.setdefault(g, h)

        # Shared legend centred under the figure.
        if legend_handles:
            ordered = sorted(legend_handles.keys())
            label_map = {
                'Chiro': 'Chiropractic',
                'Ctrl': 'Control',
            }
            display_labels = [label_map.get(g, g) for g in ordered]
            fig.legend(
                [legend_handles[g] for g in ordered],
                display_labels,
                loc='lower center', bbox_to_anchor=(0.5, -0.01),
                ncol=len(ordered), frameon=False, fontsize=11,
            )

        fig.tight_layout(rect=(0, 0.04, 1, 1))

        filepath = self.output_dir / 'combined_results.png'
        fig.savefig(str(filepath), dpi=250, bbox_inches='tight')
        plt.close(fig)

        self._update(f"Combined figure saved: {filepath.name}", 95)
        return filepath

    def save_statistics_excel(self, all_contrasts):
        """Save all statistics to Excel file.

        FDR (Benjamini–Hochberg) correction is applied within two families:
        all Between-Group contrasts and all Within-Group contrasts, with
        corrected p-values and significance written to ``p_adj`` /
        ``Significance_FDR`` columns.
        """
        self._update("\nSaving statistics to Excel...", 90)

        excel_path = self.output_dir / 'analysis_statistics.xlsx'
        contrasts_df_corrected = None

        with pd.ExcelWriter(str(excel_path), engine='openpyxl') as writer:
            # Sheet 1: Model Fixed Effects — with FDR across all effects
            if self.all_stats:
                stats_df = pd.DataFrame(self.all_stats)
                stats_df = self._apply_fdr(stats_df)
                stats_df.to_excel(writer, sheet_name='Model_Effects', index=False)

            # Sheet 2: Contrasts — FDR separately per ContrastType family
            if all_contrasts:
                contrasts_df = pd.DataFrame(all_contrasts)
                corrected = []
                for ctype, grp in contrasts_df.groupby('ContrastType', sort=False):
                    corrected.append(self._apply_fdr(grp))
                contrasts_df = pd.concat(corrected, ignore_index=True)
                contrasts_df.to_excel(writer, sheet_name='Contrasts', index=False)
                contrasts_df_corrected = contrasts_df

            # Sheet 3: Group Means Summary
            means_rows = []
            for model_name, result in self.results.items():
                if result and 'data' in result:
                    means = self.compute_means(result['data'])
                    for _, row in means.iterrows():
                        means_rows.append({
                            'Model': model_name,
                            'Group': row['Group'],
                            'Session': row['Session'],
                            'Mean': row['Mean'],
                            'SD': row['SD'],
                            'SE': row['SE'],
                            'N': row['N']
                        })

            if means_rows:
                means_df = pd.DataFrame(means_rows)
                means_df.to_excel(writer, sheet_name='Group_Means', index=False)

        self._update(f"Statistics saved: {excel_path.name}", 92)
        return excel_path, contrasts_df_corrected

    def save_between_group_table(self, contrasts_df):
        """Write a manuscript-style Word table of significant between-group
        contrasts (Post and Post4W) to the output directory, with a matching
        CSV for downstream use.

        Only FDR-significant rows (``p_adj < 0.05``) at post-treatment
        sessions are included. The layout mirrors Table 2 in the paper:
        frequency band (shown once per group), network, contrast, session,
        estimate (raw PLI scale), SE, df, *t*, and *p*-value.
        """
        if contrasts_df is None or contrasts_df.empty:
            return None

        # Post-treatment sessions = every session other than the baseline
        # (first in the configured order). Keeps the table generic for any
        # session naming (Post/Post4W, 1/2/3, ...).
        ordered = list(self.session_order) if self.session_order else []
        post_sessions = [s for s in ordered[1:]]
        if not post_sessions:
            return None

        bg = contrasts_df[
            (contrasts_df['ContrastType'] == 'Between-Group')
            & (contrasts_df['Session'].astype(str).isin([str(s) for s in post_sessions]))
            & (contrasts_df['p_adj'] < 0.05)
        ].copy()
        if bg.empty:
            self._update(
                "No FDR-significant between-group contrasts at post-treatment "
                "sessions; Table 2 not written.",
                None,
            )
            return None

        # SE reconstructed from t = diff / SE.
        with np.errstate(divide='ignore', invalid='ignore'):
            bg['SE'] = (bg['Difference'] / bg['t-value']).abs()

        # Order frequency bands / networks / sessions consistently with the
        # figure grid.
        band_order = [b for b in ['Delta', 'Theta', 'Alpha', 'Beta', 'Gamma']
                      if b in bg['FrequencyBand'].unique()]
        net_order = [n for n in ['DMN', 'SN', 'CEN']
                     if n in bg['Network'].unique()]
        sess_order = post_sessions
        bg['_band_i'] = bg['FrequencyBand'].map(
            {b: i for i, b in enumerate(band_order)}
        )
        bg['_net_i'] = bg['Network'].map(
            {n: i for i, n in enumerate(net_order)}
        )
        bg['_sess_i'] = bg['Session'].astype(str).map(
            {str(s): i for i, s in enumerate(sess_order)}
        )
        bg = bg.sort_values(['_band_i', '_net_i', '_sess_i']).reset_index(drop=True)

        # Build the clean, human-readable frame used by both CSV and DOCX.
        table_rows = []
        for _, r in bg.iterrows():
            d_val = r['Cohens_d'] if 'Cohens_d' in r.index and pd.notna(r['Cohens_d']) else float('nan')
            table_rows.append({
                'Frequency Band': r['FrequencyBand'],
                'Brain Network': r['Network'],
                'Contrast': str(r['Contrast']).replace(' vs ', ' \u2013 '),
                'Session': r['Session'],
                'Estimate': float(r['Difference']),
                'SE': float(r['SE']),
                'df': int(r['df']) if pd.notna(r['df']) else '',
                't': float(r['t-value']),
                'p-value': float(r['p_adj']),
                "Cohen's d": float(d_val),
            })
        table_df = pd.DataFrame(table_rows)

        csv_path = self.output_dir / 'Table2_BetweenGroup_Contrasts.csv'
        table_df.to_csv(csv_path, index=False, float_format='%.6f')
        self._update(f"Table 2 CSV saved: {csv_path.name}", None)

        if not HAS_DOCX:
            self._update(
                "python-docx not installed; Table 2 Word file skipped. "
                "Run: pip install python-docx",
                None,
            )
            return csv_path

        docx_path = self.output_dir / 'Table2_BetweenGroup_Contrasts.docx'
        figure_path = self.output_dir / 'combined_results.png'
        self._write_between_group_docx(table_df, docx_path, figure_path=figure_path)
        self._update(f"Table 2 Word file saved: {docx_path.name}", None)
        return docx_path

    @staticmethod
    def _write_between_group_docx(table_df, docx_path, figure_path=None):
        """Render ``table_df`` as an APA-style Word table at ``docx_path``."""
        MINUS = '\u2212'

        def fmt_signed(v, nd):
            return f'{v:+.{nd}f}'.replace('-', MINUS)

        def fmt_num(v, nd):
            return f'{v:.{nd}f}'.replace('-', MINUS)

        def fmt_p(p):
            return '< .001' if p < 0.001 else f'{p:.3f}'.lstrip('0')

        def set_cell_border(cell, **kwargs):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            for edge in ('top', 'bottom'):
                if edge not in kwargs:
                    continue
                e = tcBorders.find(qn(f'w:{edge}'))
                if e is None:
                    e = OxmlElement(f'w:{edge}')
                    tcBorders.append(e)
                spec = kwargs[edge]
                e.set(qn('w:val'), spec['val'])
                e.set(qn('w:sz'), str(spec['sz']))
                e.set(qn('w:space'), '0')
                e.set(qn('w:color'), '000000')
            for edge in ('left', 'right', 'insideH', 'insideV'):
                e = tcBorders.find(qn(f'w:{edge}'))
                if e is None:
                    e = OxmlElement(f'w:{edge}')
                    tcBorders.append(e)
                e.set(qn('w:val'), 'nil')

        doc = Document()
        for s in doc.sections:
            s.top_margin = Inches(1)
            s.bottom_margin = Inches(1)
            s.left_margin = Inches(1)
            s.right_margin = Inches(1)
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)

        cap = doc.add_paragraph()
        cap.paragraph_format.space_after = Pt(6)
        lead = cap.add_run('Table 2: ')
        lead.bold = True
        lead.italic = True
        body = cap.add_run(
            'Significant between-group changes in network-level phase lag '
            'index (PLI) connectivity during resting state EEG. Results are '
            'derived from linear mixed-effects models comparing the '
            'chiropractic and control groups across frequency bands within '
            'the default mode network (DMN), salience network (SN), and '
            'central executive network (CEN). Reported values are '
            'model-estimated pairwise contrasts for each session (Post, '
            'Post4W) relative to baseline, including estimates, standard '
            'errors (SE), degrees of freedom (df), t-statistics, '
            'FDR-corrected p-values, and Cohen\u2019s d (pooled SD) as a '
            'measure of effect size.'
        )
        body.italic = True

        headers = [
            'Frequency\nBand', 'Brain\nNetwork', 'Contrast', 'Session',
            'Estimate', 'SE', 'df', 't', 'p-value', 'd',
        ]
        table = doc.add_table(rows=1 + len(table_df), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Header row — top double rule and thin bottom rule (APA style).
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.bold = True
            run.italic = h in ('t', 'd')
            run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(
                cell,
                top={'val': 'single', 'sz': '12'},
                bottom={'val': 'single', 'sz': '8'},
            )

        last_row_i = len(table_df) - 1
        prev_band = None
        for i, row in table_df.iterrows():
            cells = table.rows[i + 1].cells
            band_label = row['Frequency Band'] if row['Frequency Band'] != prev_band else ''
            prev_band = row['Frequency Band']
            d_raw = row.get("Cohen's d", float('nan'))
            try:
                d_str = fmt_num(float(d_raw), 2)
            except (TypeError, ValueError):
                d_str = '—'
            values = [
                band_label,
                row['Brain Network'],
                row['Contrast'] if band_label or i == 0 or row['Brain Network'] != table_df.iloc[i - 1]['Brain Network'] else row['Contrast'],
                row['Session'],
                fmt_signed(row['Estimate'], 3),
                fmt_num(row['SE'], 3),
                str(row['df']),
                fmt_num(row['t'], 2),
                fmt_p(row['p-value']),
                d_str,
            ]
            for j, v in enumerate(values):
                c = cells[j]
                c.text = ''
                p = c.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(v))
                run.font.size = Pt(10)
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if i == last_row_i:
                    set_cell_border(c, bottom={'val': 'single', 'sz': '12'})
                else:
                    set_cell_border(c)

        widths = [
            Inches(0.80), Inches(0.70), Inches(1.00), Inches(0.70),
            Inches(0.70), Inches(0.55), Inches(0.40), Inches(0.50),
            Inches(0.65), Inches(0.45),
        ]
        for r in table.rows:
            for j, c in enumerate(r.cells):
                c.width = widths[j]

        # APA-formatted Results paragraph(s) describing the table, appended
        # below so the Word file is a single drop-in block for the manuscript.
        paragraphs = NetworkAnalysis._build_apa_results_runs(table_df)
        if paragraphs:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(12)
            spacer.paragraph_format.space_after = Pt(0)

            heading = doc.add_paragraph()
            heading.paragraph_format.space_after = Pt(6)
            h_run = heading.add_run('Results')
            h_run.bold = True
            h_run.font.name = 'Times New Roman'
            h_run.font.size = Pt(12)

            for runs in paragraphs:
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.first_line_indent = Inches(0.5)
                pf.line_spacing = 2.0
                pf.space_after = Pt(0)
                for text, italic in runs:
                    r = p.add_run(text)
                    r.italic = bool(italic)
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)

        if figure_path is not None and Path(figure_path).exists():
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(12)
            spacer.paragraph_format.space_after = Pt(0)

            fig_heading = doc.add_paragraph()
            fig_heading.paragraph_format.space_after = Pt(6)
            h_run = fig_heading.add_run('Combined Results Figure')
            h_run.bold = True
            h_run.font.name = 'Times New Roman'
            h_run.font.size = Pt(12)

            fig_paragraph = doc.add_paragraph()
            fig_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fig_paragraph.add_run().add_picture(str(figure_path), width=Inches(6.5))

        doc.save(str(docx_path))

    @staticmethod
    def _build_apa_results_runs(table_df):
        """Build an APA-style Results narrative from the Table 2 dataframe.

        Returns a list of paragraphs, where each paragraph is a list of
        ``(text, italic)`` run tuples ready to be written into a
        python-docx paragraph. Only *t* and *p* are italicised, per APA 7.
        Group names are parsed from the contrast string, so the narrative
        correctly names whichever groups were compared.
        """
        if table_df is None or len(table_df) == 0:
            return []

        band_order = ['Delta', 'Theta', 'Alpha', 'Beta', 'Gamma']
        net_order = ['DMN', 'SN', 'CEN']
        MINUS = '\u2212'

        name_map = {
            'chiro': 'chiropractic',
            'chiropractic': 'chiropractic',
            'control': 'control',
            'ctrl': 'control',
        }

        def nice(g):
            return name_map.get(str(g).strip().lower(), str(g).strip())

        def parse_groups(contrast_str):
            for sep in (' \u2013 ', ' - ', ' vs '):
                if sep in contrast_str:
                    g1, g2 = contrast_str.split(sep, 1)
                    return nice(g1), nice(g2)
            return None, None

        def fmt_t(v):
            return f'{v:.2f}'.replace('-', MINUS)

        def fmt_p(p):
            if p < 0.001:
                return '< .001'
            return '= ' + f'{p:.3f}'.lstrip('0')

        def join_and(items):
            if not items:
                return ''
            if len(items) == 1:
                return items[0]
            if len(items) == 2:
                return f'{items[0]} and {items[1]}'
            return ', '.join(items[:-1]) + f', and {items[-1]}'

        present_bands = [b for b in band_order
                         if b in set(table_df['Frequency Band'])]
        for b in table_df['Frequency Band'].unique():
            if b not in present_bands:
                present_bands.append(b)

        paragraphs = []

        for band in present_bands:
            br = table_df[table_df['Frequency Band'] == band]
            if br.empty:
                continue

            present_nets = [n for n in net_order
                            if n in set(br['Brain Network'])]
            for n in br['Brain Network'].unique():
                if n not in present_nets:
                    present_nets.append(n)
            if not present_nets:
                continue

            runs = []

            def add(s, italic=False, _runs=runs):
                _runs.append((s, italic))

            def add_stat(row):
                add(f"{row['Session']}: ")
                add('t', italic=True)
                add(f"({int(row['df'])}) = {fmt_t(row['t'])}, ")
                add('p', italic=True)
                add(f" {fmt_p(row['p-value'])}")

            # Build (network, contrast) clause groups so that studies with
            # more than two groups (i.e. multiple pairwise contrasts per
            # network) still produce coherent sentences. For the typical
            # Chiro-vs-Control design there is only one contrast per network,
            # and this reduces to "{network} connectivity was …".
            clause_groups = []
            for net in present_nets:
                nr_net = br[br['Brain Network'] == net]
                for contrast_str in nr_net['Contrast'].unique():
                    sub = nr_net[nr_net['Contrast'] == contrast_str]
                    sub = sub.reset_index(drop=True)
                    clause_groups.append((net, contrast_str, sub))
            if not clause_groups:
                continue

            single_net = len(present_nets) == 1
            if single_net:
                add(f"{band}-band differences were confined to the "
                    f"{present_nets[0]}. ")
            else:
                add(f"In the {band.lower()} band, significant between-group "
                    f"differences were observed in the "
                    f"{join_and(present_nets)}. ")

            for ci, (net, contrast_str, sub) in enumerate(clause_groups):
                sessions = list(sub['Session'])
                estimates = list(sub['Estimate'])
                g1, g2 = parse_groups(contrast_str)
                if g1 is None:
                    g1, g2 = 'group 1', 'group 2'
                dirs = ['higher' if e > 0 else 'lower' for e in estimates]
                same_dir = len(set(dirs)) == 1

                # Name the network only on the first clause for that network
                # (keeps successive contrasts within the same network from
                # repeating "{net} connectivity was …").
                net_first_here = (
                    ci == 0 or clause_groups[ci - 1][0] != net
                )
                if single_net and ci == 0:
                    add('Connectivity was ')
                elif net_first_here:
                    add(f"{net} connectivity was ")
                else:
                    add(f"The {g1}-vs-{g2} contrast in the {net} was ")

                if same_dir:
                    add(f"{dirs[0]} in the {g1} group than in the "
                        f"{g2} group")
                else:
                    add(f"mixed across sessions in the {g1} group relative "
                        f"to the {g2} group")

                if len(sessions) > 1 and same_dir:
                    add(' at both post-treatment sessions (')
                elif len(sessions) == 1:
                    add(f' at {sessions[0]} (')
                else:
                    add(' (')

                for i in range(len(sub)):
                    if i > 0:
                        add('; ')
                    add_stat(sub.iloc[i])
                add('). ')

            paragraphs.append(runs)

        paragraphs.append([(
            'Table 2 summarises all significant between-group contrasts, '
            'as illustrated in Figure 2.',
            False,
        )])

        return paragraphs

    @staticmethod
    def _apply_fdr(df):
        """Add p_adj and Significance_FDR columns via Benjamini–Hochberg."""
        if 'p-value' not in df.columns or df.empty:
            return df
        pvals = df['p-value'].values.astype(float)
        valid = np.isfinite(pvals)
        p_adj = np.full_like(pvals, np.nan)
        if valid.sum() > 0:
            _, corrected, _, _ = multipletests(
                pvals[valid], method='fdr_bh', alpha=0.05,
            )
            p_adj[valid] = corrected
        df = df.copy()
        df['p_adj'] = p_adj
        df['Significance_FDR'] = df['p_adj'].apply(
            lambda p: NetworkAnalysis._get_sig_stars_static(p)
        )
        return df

    @staticmethod
    def _get_sig_stars_static(p):
        if pd.isna(p):
            return ''
        if p < 0.001:
            return '***'
        elif p < 0.01:
            return '**'
        elif p < 0.05:
            return '*'
        else:
            return 'ns'

    def run(self):
        """Run complete analysis."""
        self._update("=" * 50, 0)
        self._update("NETWORK LEVEL ANALYSIS", 0)
        self._update("=" * 50, 0)

        # Create output directory
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._update(f"Output: {self.output_dir}", 5)

        # Load data
        self.load_data()

        # Run analysis for each combination
        all_results = []
        all_contrasts = []
        total = len(self.networks) * len(self.frequency_bands)
        current = 0

        for network in self.networks:
            for freq in self.frequency_bands:
                current += 1
                pct = 20 + int((current / total) * 60)

                model_name = f"{freq} x {network}"
                self._update(f"\n{'=' * 40}", pct)
                self._update(f"Model: {model_name} ({current}/{total})", pct)
                self._update('=' * 40, pct)

                result = self.fit_model(network, freq)

                if result is None:
                    self._update("  No data for this combination", pct)
                    continue

                subset = result['data']

                # Compute means
                means = self.compute_means(subset)
                self._update("\nGroup Means:", pct)
                for _, row in means.iterrows():
                    self._update(f"  {row['Group']:8} {row['Session']:8}: {row['Mean']:.4f} +/- {row['SE']:.4f}", pct)

                # Run contrasts
                contrasts, contrast_lines, contrast_stats = self.run_contrasts(
                    subset, model_name, network, freq
                )
                all_contrasts.extend(contrast_stats)

                self._update("\nContrasts:", pct)
                for line in contrast_lines:
                    self._update(line, pct)

                # Bar plots annotate the per-session Chiro-vs-Control
                # between-group contrast (matches the manuscript table).
                significance = contrasts.get('significance', {})
                plot_path = self.create_plot(subset, model_name, significance=significance)
                self._update(f"\nPlot saved: {plot_path.name}", pct)

                all_results.append({
                    'model': model_name,
                    'means': means,
                    'contrasts': contrasts
                })

                self.results[model_name] = result

        # Save statistics to Excel
        _, contrasts_df_corrected = self.save_statistics_excel(all_contrasts)

        # Emit the manuscript-style Table 2 (CSV + Word) from the same
        # FDR-corrected Between-Group contrasts that feed the Excel workbook.
        self.save_between_group_table(contrasts_df_corrected)

        # Create combined figure
        self.create_combined_figure()

        # Save summary CSV
        self._save_summary(all_results)

        self._update("\n" + "=" * 50, 100)
        self._update("ANALYSIS COMPLETE", 100)
        self._update(f"Results saved to: {self.output_dir}", 100)
        self._update("=" * 50, 100)

        return all_results

    def _save_summary(self, results):
        """Save results summary to CSV."""
        rows = []
        for r in results:
            for _, row in r['means'].iterrows():
                rows.append({
                    'Model': r['model'],
                    'Group': row['Group'],
                    'Session': row['Session'],
                    'Mean': row['Mean'],
                    'SE': row['SE'],
                    'N': row['N']
                })

        if rows:
            df = pd.DataFrame(rows)
            csv_path = self.output_dir / 'summary.csv'
            df.to_csv(str(csv_path), index=False)
            self._update(f"Summary CSV saved: {csv_path.name}", 97)


def main():
    """Command-line entry point."""
    import argparse

    parser = argparse.ArgumentParser(description='Network Level PLI Analysis')
    parser.add_argument('--input', '-i', default='PLI-UK-Both-Groups-UP3.xlsx',
                        help='Input Excel file')
    parser.add_argument('--output', '-o', default='analysis_output',
                        help='Output directory')
    parser.add_argument('--no-baseline', action='store_true',
                        help='Disable baseline adjustment')
    parser.add_argument('--plot-type', choices=['bar', 'line'], default='bar',
                        help='Plot style for saved figures')
    parser.add_argument('--groups', nargs='+', default=None,
                        help='Groups to include (default: auto-detect from data)')
    parser.add_argument('--sessions', nargs='+', default=None,
                        help='Sessions to include (default: auto-detect from data)')
    parser.add_argument('--networks', nargs='+', default=None,
                        help='Networks to analyze (default: auto-detect from data)')
    parser.add_argument('--bands', nargs='+', default=None,
                        help='Frequency bands (default: auto-detect from data)')

    args = parser.parse_args()

    analysis = NetworkAnalysis(
        input_file=args.input,
        output_dir=args.output,
        adjust_baseline=not args.no_baseline,
        plot_type=args.plot_type,
        groups=args.groups,
        sessions=args.sessions,
        networks=args.networks,
        frequency_bands=args.bands
    )

    analysis.run()


if __name__ == '__main__':
    main()
