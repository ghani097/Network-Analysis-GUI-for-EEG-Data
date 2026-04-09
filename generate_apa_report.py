"""
Generate APA-style publication-ready Word document for PLI Network Analysis.
Focused contrasts:
  1. Between-group: Group A vs Group B (all sessions)
  2. Within-group: pre1 vs post1  (Lower Cervical adjustment)
  3. Within-group: pre2 vs post2  (Upper Cervical adjustment)
"""

import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── paths ──────────────────────────────────────────────────────────────────────
BASE    = Path("E:/GIT_HUB_MAIN/PLI-Network-Analysis-GUI")
OUT_DIR = BASE / "analysis_output"
XLSX    = OUT_DIR / "analysis_statistics.xlsx"
FIG     = OUT_DIR / "combined_results.png"
OUTPUT  = OUT_DIR / "Network_Analysis_Results_APA.docx"

# ── load data ──────────────────────────────────────────────────────────────────
gm        = pd.read_excel(XLSX, sheet_name="Group_Means")
contrasts = pd.read_excel(XLSX, sheet_name="Contrasts")

# ── derive lookup for Table 1 ─────────────────────────────────────────────────
gm["FreqBand"] = gm["Model"].str.split(" x ").str[0]
gm["Net"]      = gm["Model"].str.split(" x ").str[1]
lookup = {
    (r.FreqBand, r.Net, r.Group, r.Session): (r.Mean, r.SE)
    for _, r in gm.iterrows()
}

# ── filter contrasts to our three categories ──────────────────────────────────
between = contrasts[contrasts["ContrastType"] == "Between-Group"].copy()
lc      = contrasts[
    (contrasts["ContrastType"] == "Within-Group") &
    (contrasts["Contrast"] == "pre1 vs post1")
].copy()
uc      = contrasts[
    (contrasts["ContrastType"] == "Within-Group") &
    (contrasts["Contrast"] == "pre2 vs post2")
].copy()

# Significant subsets
sig_between = between[between["p-value"] < 0.05].copy()
sig_lc      = lc[lc["p-value"] < 0.05].copy()
sig_uc      = uc[uc["p-value"] < 0.05].copy()

# ── APA border helpers ─────────────────────────────────────────────────────────
def _apply_border(cell, edge, sz=12):
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcB   = OxmlElement("w:tcBorders")
    tag   = OxmlElement(f"w:{edge}")
    tag.set(qn("w:val"),   "single")
    tag.set(qn("w:sz"),    str(sz))
    tag.set(qn("w:space"), "0")
    tag.set(qn("w:color"), "000000")
    tcB.append(tag)
    tcPr.append(tcB)

def clear_all_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcB  = OxmlElement("w:tcBorders")
            for edge in ["top","bottom","left","right","insideH","insideV"]:
                tag = OxmlElement(f"w:{edge}")
                tag.set(qn("w:val"),   "none")
                tag.set(qn("w:sz"),    "0")
                tag.set(qn("w:space"), "0")
                tag.set(qn("w:color"), "auto")
                tcB.append(tag)
            tcPr.append(tcB)

def top_border(cell, sz=18):    _apply_border(cell, "top",    sz)
def bottom_border(cell, sz=18): _apply_border(cell, "bottom", sz)

# ── APA run / paragraph helpers ────────────────────────────────────────────────
def _new_para(doc, indent=False, centered=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing       = Pt(24)
    p.paragraph_format.line_spacing_rule  = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.space_before       = Pt(0)
    p.paragraph_format.space_after        = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Pt(36)
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def _run(para, text, bold=False, italic=False, size=12):
    r = para.add_run(text)
    r.bold        = bold
    r.italic      = italic
    r.font.size   = Pt(size)
    r.font.name   = "Times New Roman"
    return r

def _note_para(doc, lead="Note. ", body=""):
    p = _new_para(doc)
    p.paragraph_format.line_spacing      = Pt(18)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    _run(p, lead, bold=True, size=10)
    _run(p, body, size=10)
    return p

def _tbl_cell_fmt(table):
    """Apply compact 10-pt formatting to all table cells."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before      = Pt(2)
                para.paragraph_format.space_after       = Pt(2)
                para.paragraph_format.line_spacing      = Pt(14)
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"

def _set_col_widths(table, widths):
    for i, w in enumerate(widths):
        for row in table.rows:
            try:
                row.cells[i].width = w
            except Exception:
                pass

# ── build document ─────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing      = Pt(24)
pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
pf.space_before      = Pt(0)
pf.space_after       = Pt(0)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION HEADING
# ═══════════════════════════════════════════════════════════════════════════════
p = _new_para(doc, centered=True)
_run(p, "Results", bold=True)

p = _new_para(doc)
_run(p, "Network-Level EEG Connectivity", bold=True)


# ═══════════════════════════════════════════════════════════════════════════════
#  OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
p = _new_para(doc, indent=True)
_run(p,
     "Phase lag index (PLI) was computed as a measure of mean functional "
     "connectivity within three canonical resting-state networks—the central "
     "executive network (CEN), the default mode network (DMN), and the salience "
     "network (SN)—across five frequency bands (delta, theta, alpha, beta, and "
     "gamma) at four assessment points: baseline prior to lower cervical adjustment "
     "(pre1), immediately following lower cervical adjustment (post1), baseline "
     "prior to upper cervical adjustment (pre2), and immediately following upper "
     "cervical adjustment (post2). Group A (")
_run(p, "n", italic=True)
_run(p, " = 2) and Group B (")
_run(p, "n", italic=True)
_run(p,
     " = 7) were compared using separate linear mixed-effects models for each of "
     "the 15 frequency-band × network combinations, with Group, Session, and their "
     "interaction as fixed effects and participant as a random intercept. Three "
     "pre-specified contrast families were evaluated: (1) between-group differences "
     "(Group A vs. Group B) at each session, (2) within-group change following lower "
     "cervical adjustment (pre1 vs. post1), and (3) within-group change following "
     "upper cervical adjustment (pre2 vs. post2). Descriptive statistics are "
     "presented in Table 1 and mean PLI trajectories are illustrated in Figure 1. "
     "Statistically significant contrasts are summarised in Table 2.")


# ═══════════════════════════════════════════════════════════════════════════════
#  HEADING 3 — BETWEEN-GROUP DIFFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
p = _new_para(doc, indent=True)
r = _run(p, "Group A vs. Group B Differences.", bold=True, italic=True)

p = _new_para(doc, indent=True)
_run(p,
     "Five significant between-group contrasts emerged across the 15 models. "
     "At the second baseline (pre2), Group A exhibited significantly higher delta "
     "PLI than Group B within both the CEN (")
_run(p, "M", italic=True)
_run(p, " difference = 0.051, ")
_run(p, "t", italic=True)
_run(p, "(13) = 3.88, ")
_run(p, "p", italic=True)
_run(p,
     " = .006) and the DMN (")
_run(p, "M", italic=True)
_run(p, " difference = 0.047, ")
_run(p, "t", italic=True)
_run(p, "(13) = 3.66, ")
_run(p, "p", italic=True)
_run(p,
     " = .008), indicating that Group A entered the upper cervical session with "
     "markedly elevated low-frequency functional coupling in both task-control "
     "and default-mode circuits relative to Group B.")

p = _new_para(doc, indent=True)
_run(p,
     "In the salience network, Group A showed substantially higher alpha PLI than "
     "Group B at both post1 (")
_run(p, "M", italic=True)
_run(p, " difference = 0.107, ")
_run(p, "t", italic=True)
_run(p, "(13) = 2.62, ")
_run(p, "p", italic=True)
_run(p, " = .034) and post2 (")
_run(p, "M", italic=True)
_run(p, " difference = 0.114, ")
_run(p, "t", italic=True)
_run(p, "(13) = 3.19, ")
_run(p, "p", italic=True)
_run(p,
     " = .015), and higher delta PLI at post1 (")
_run(p, "M", italic=True)
_run(p, " difference = 0.054, ")
_run(p, "t", italic=True)
_run(p, "(13) = 3.18, ")
_run(p, "p", italic=True)
_run(p,
     " = .019). Collectively, these between-group contrasts indicate that Group A "
     "maintained persistently elevated salience network connectivity—across both "
     "alpha and delta bands—following each adjustment session, whereas Group B "
     "showed a more attenuated post-adjustment SN response.")


# ═══════════════════════════════════════════════════════════════════════════════
#  HEADING 3 — LOWER CERVICAL ADJUSTMENT  (pre1 → post1)
# ═══════════════════════════════════════════════════════════════════════════════
p = _new_para(doc, indent=True)
_run(p, "Lower Cervical Adjustment Effects (Pre1 to Post1).", bold=True, italic=True)

p = _new_para(doc, indent=True)
_run(p,
     "Examining within-group change from pre1 to post1 across both groups and all "
     "15 frequency-band × network combinations, one significant effect emerged. "
     "In Group A, beta-band PLI within the CEN decreased from pre1 to post1 (")
_run(p, "M", italic=True)
_run(p, " difference = 0.021, ")
_run(p, "t", italic=True)
_run(p, "(13) = 4.79, ")
_run(p, "p", italic=True)
_run(p,
     " = .041), reflecting a post-adjustment suppression of beta-band functional "
     "coupling within the central executive network. No significant pre1-to-post1 "
     "changes were observed for Group B in any frequency band or network, nor were "
     "significant effects detected in the DMN or SN for either group.")


# ═══════════════════════════════════════════════════════════════════════════════
#  HEADING 3 — UPPER CERVICAL ADJUSTMENT  (pre2 → post2)
# ═══════════════════════════════════════════════════════════════════════════════
p = _new_para(doc, indent=True)
_run(p, "Upper Cervical Adjustment Effects (Pre2 to Post2).", bold=True, italic=True)

p = _new_para(doc, indent=True)
_run(p,
     "Within-group change from pre2 to post2 yielded one significant contrast. "
     "In Group A, delta-band PLI within the SN increased significantly from pre2 "
     "to post2 (")
_run(p, "M", italic=True)
_run(p, " difference = 0.062, ")
_run(p, "t", italic=True)
_run(p, "(13) = 4.88, ")
_run(p, "p", italic=True)
_run(p,
     " = .040), indicating a strengthening of low-frequency salience network "
     "connectivity following the upper cervical adjustment. No significant "
     "pre2-to-post2 changes were observed for Group B in any frequency band or "
     "network, and no other CEN or DMN effects reached significance for either group.")


# ═══════════════════════════════════════════════════════════════════════════════
#  SUMMARY PARAGRAPH
# ═══════════════════════════════════════════════════════════════════════════════
p = _new_para(doc, indent=True)
_run(p,
     "In summary, statistically significant connectivity differences were "
     "concentrated in the delta and alpha frequency bands and were most prominent "
     "within the SN, with additional effects in the CEN. Group A consistently "
     "exhibited higher post-adjustment SN connectivity than Group B (alpha and "
     "delta bands), and showed a transient suppression of CEN beta connectivity "
     "after the lower cervical adjustment and an increase in SN delta connectivity "
     "after the upper cervical adjustment. Group B showed no significant within-group "
     "changes under either adjustment condition. These patterns suggest differential "
     "network reconfiguration between groups in response to cervical spinal adjustment, "
     "with Group A demonstrating greater post-adjustment modulation of salience and "
     "executive network dynamics. All significant contrasts are presented in Table 2 "
     "and visualised in Figure 1.")


# ═══════════════════════════════════════════════════════════════════════════════
#  TABLE 1  –  Descriptive Statistics
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()

p = _new_para(doc)
_run(p, "Table 1", bold=True)

p = _new_para(doc)
_run(p,
     "Mean PLI Values (Standard Error) by Group, Frequency Band, "
     "and Network Across Sessions")

BANDS       = ["delta", "theta", "alpha", "beta", "gamma"]
NETWORKS    = ["CEN", "DMN", "SN"]
SESSIONS    = ["pre1", "post1", "pre2", "post2"]
SESS_LABELS = ["Pre1\n(Baseline LC)", "Post1\n(Post LC)", "Pre2\n(Baseline UC)", "Post2\n(Post UC)"]
SESS_LABELS_PLAIN = ["Pre1 (Baseline LC)", "Post1 (Post LC)", "Pre2 (Baseline UC)", "Post2 (Post UC)"]

# Cells that have a significant between-group contrast
sig_bg_keys = set(
    zip(sig_between["FrequencyBand"], sig_between["Network"], sig_between["Session"])
)

n_data = len(BANDS) * len(NETWORKS)
n_rows = 2 + n_data           # header-group row + header-session row + data
n_cols = 10                   # Band | Net | GA×4 | GB×4

tbl = doc.add_table(rows=n_rows, cols=n_cols)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
clear_all_borders(tbl)

# ── Row 0 : Group-level headers ───────────────────────────────────────────────
r0 = tbl.rows[0]
for c in r0.cells:
    top_border(c, sz=18)

r0.cells[0].merge(r0.cells[1])
r0.cells[0].paragraphs[0].clear()

ga = r0.cells[2]; ga.merge(r0.cells[5])
p = ga.paragraphs[0]; p.clear()
_run(p, "Group A", bold=True, size=10); p.alignment = WD_ALIGN_PARAGRAPH.CENTER

gb = r0.cells[6]; gb.merge(r0.cells[9])
p = gb.paragraphs[0]; p.clear()
_run(p, "Group B", bold=True, size=10); p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for c in r0.cells:
    bottom_border(c, sz=6)

# ── Row 1 : Session sub-headers ───────────────────────────────────────────────
r1 = tbl.rows[1]
r1.cells[0].paragraphs[0].clear(); _run(r1.cells[0].paragraphs[0], "Band", bold=True, size=10)
r1.cells[1].paragraphs[0].clear(); _run(r1.cells[1].paragraphs[0], "Network", bold=True, size=10)

for gi, goff in enumerate([2, 6]):
    for si, label in enumerate(SESS_LABELS_PLAIN):
        c = r1.cells[goff + si]
        p = c.paragraphs[0]; p.clear()
        _run(p, label, bold=True, size=10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for c in r1.cells:
    bottom_border(c, sz=6)

# ── Data rows ─────────────────────────────────────────────────────────────────
for bi, band in enumerate(BANDS):
    for ni, net in enumerate(NETWORKS):
        ri  = 2 + bi * len(NETWORKS) + ni
        row = tbl.rows[ri]

        row.cells[0].paragraphs[0].clear()
        _run(row.cells[0].paragraphs[0], band.capitalize() if ni == 0 else "", size=10)
        row.cells[1].paragraphs[0].clear()
        _run(row.cells[1].paragraphs[0], net, size=10)

        for gi, group in enumerate(["Group A", "Group B"]):
            goff = 2 + gi * 4
            for si, sess in enumerate(SESSIONS):
                val = lookup.get((band, net, group, sess))
                c   = row.cells[goff + si]
                p   = c.paragraphs[0]; p.clear()
                if val:
                    m, se   = val
                    is_sig  = (band, net, sess) in sig_bg_keys
                    suffix  = "\u2020" if is_sig else ""        # dagger for sig between-group
                    _run(p, f"{m:.3f} ({se:.3f}){suffix}", size=10)
                else:
                    _run(p, "\u2014", size=10)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Bottom rule on last data row
for c in tbl.rows[-1].cells:
    bottom_border(c, sz=18)

# Column widths
_set_col_widths(tbl, [
    Inches(0.60), Inches(0.52),
    Inches(0.82), Inches(0.82), Inches(0.82), Inches(0.82),
    Inches(0.82), Inches(0.82), Inches(0.82), Inches(0.82),
])
_tbl_cell_fmt(tbl)

# ── Table 1 note ──────────────────────────────────────────────────────────────
p = _new_para(doc)
p.paragraph_format.line_spacing      = Pt(18)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
_run(p, "Note. ", bold=True, size=10)
_run(p,
     "Values are mean PLI (standard error). "
     "PLI = Phase Lag Index; CEN = Central Executive Network; "
     "DMN = Default Mode Network; SN = Salience Network; "
     "LC = Lower Cervical adjustment; UC = Upper Cervical adjustment. "
     "Group A: ", size=10)
_run(p, "n", italic=True, size=10)
_run(p, " = 2; Group B: ", size=10)
_run(p, "n", italic=True, size=10)
_run(p,
     " = 7. \u2020Between-group contrast (Group A vs. Group B) significant at ",
     size=10)
_run(p, "p", italic=True, size=10)
_run(p, " < .05.", size=10)


# ═══════════════════════════════════════════════════════════════════════════════
#  TABLE 2  –  Significant Contrasts
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()

p = _new_para(doc)
_run(p, "Table 2", bold=True)

p = _new_para(doc)
_run(p,
     "Statistically Significant Post-Hoc Contrasts (p < .05) for Group "
     "Differences, Lower Cervical Adjustment, and Upper Cervical Adjustment")

# Build rows for each category
def _between_row(r):
    return {
        "Category": "Group A vs. B",
        "Network":  r["Network"],
        "Band":     r["FrequencyBand"].capitalize(),
        "Session":  r["Session"],
        "Group":    "—",
        "Direction":"Group A > Group B",
        "DeltaM":   f"{r['Difference']:.3f}",
        "t":        f"{r['t-value']:.2f}",
        "p":        f"{r['p-value']:.3f}{r['Significance']}",
    }

def _within_row(r, label):
    diff  = r["Difference"]
    parts = r["Contrast"].split(" vs ")
    s1, s2 = parts[0].strip(), parts[1].strip()
    direction = f"{s1} > {s2}" if diff > 0 else f"{s2} > {s1}"
    return {
        "Category": label,
        "Network":  r["Network"],
        "Band":     r["FrequencyBand"].capitalize(),
        "Session":  r["Contrast"],
        "Group":    r["Group"],
        "Direction": direction,
        "DeltaM":   f"{abs(diff):.3f}",
        "t":        f"{abs(r['t-value']):.2f}",
        "p":        f"{r['p-value']:.3f}{r['Significance']}",
    }

rows_t2 = (
    [_between_row(r) for _, r in sig_between.iterrows()] +
    [_within_row(r, "Lower Cervical\n(Pre1\u2192Post1)") for _, r in sig_lc.iterrows()] +
    [_within_row(r, "Upper Cervical\n(Pre2\u2192Post2)") for _, r in sig_uc.iterrows()]
)

HEADERS = ["Contrast Type", "Network", "Band", "Session / Contrast",
           "Group", "Direction", "|ΔM|", "t", "p"]
n2c = len(HEADERS)
n2r = 1 + len(rows_t2)

tbl2 = doc.add_table(rows=n2r, cols=n2c)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.style = "Table Grid"
clear_all_borders(tbl2)

# Header row
hr = tbl2.rows[0]
for c in hr.cells:
    top_border(c, sz=18)
    bottom_border(c, sz=6)

for ci, h in enumerate(HEADERS):
    cell = hr.cells[ci]
    p    = cell.paragraphs[0]; p.clear()
    # Make t and p italic
    if h in ("t", "p"):
        _run(p, h, bold=True, italic=True, size=10)
    elif h == "|ΔM|":
        _run(p, "|Δ", bold=True, size=10)
        _run(p, "M", bold=True, italic=True, size=10)
        _run(p, "|", bold=True, size=10)
    else:
        _run(p, h, bold=True, size=10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data rows
FIELD_ORDER = ["Category","Network","Band","Session","Group","Direction","DeltaM","t","p"]
CENTER_COLS  = {1, 2, 6, 7, 8}   # Network, Band, |ΔM|, t, p

for ri, rd in enumerate(rows_t2):
    row = tbl2.rows[ri + 1]
    for ci, key in enumerate(FIELD_ORDER):
        c = row.cells[ci]
        p = c.paragraphs[0]; p.clear()
        _run(p, str(rd[key]), size=10)
        if ci in CENTER_COLS:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Bottom rule
for c in tbl2.rows[-1].cells:
    bottom_border(c, sz=18)

# Column widths
_set_col_widths(tbl2, [
    Inches(0.90), Inches(0.55), Inches(0.50),
    Inches(0.85), Inches(0.70), Inches(1.00),
    Inches(0.50), Inches(0.50), Inches(0.60),
])
_tbl_cell_fmt(tbl2)

# ── Table 2 note ──────────────────────────────────────────────────────────────
p = _new_para(doc)
p.paragraph_format.line_spacing      = Pt(18)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
_run(p, "Note. ", bold=True, size=10)
_run(p, "|Δ", size=10)
_run(p, "M", italic=True, size=10)
_run(p,
     "| = absolute mean difference. Between-group contrasts compare Group A "
     "versus Group B at the specified session. Within-group contrasts compare "
     "connectivity before and after the respective adjustment within each group. "
     "LC = Lower Cervical; UC = Upper Cervical. *",
     size=10)
_run(p, "p", italic=True, size=10)
_run(p, " < .05; **", size=10)
_run(p, "p", italic=True, size=10)
_run(p, " < .01.", size=10)


# ═══════════════════════════════════════════════════════════════════════════════
#  FIGURE 1
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

p = _new_para(doc)
_run(p, "Figure 1", bold=True)

p = _new_para(doc)
_run(p,
     "Mean PLI Trajectories Across Sessions for Each Frequency-Band × "
     "Network Combination")

doc.add_picture(str(FIG), width=Inches(6.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

p = _new_para(doc)
p.paragraph_format.line_spacing      = Pt(18)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
_run(p, "Note. ", bold=True, size=10)
_run(p,
     "Mean Phase Lag Index (PLI) for Group A (orange; ",
     size=10)
_run(p, "n", italic=True, size=10)
_run(p,
     " = 2) and Group B (blue; ",
     size=10)
_run(p, "n", italic=True, size=10)
_run(p,
     " = 7) across four assessment points: Pre1 (baseline before lower cervical [LC] "
     "adjustment), Post1 (immediately after LC adjustment), Pre2 (baseline before "
     "upper cervical [UC] adjustment), and Post2 (immediately after UC adjustment). "
     "Error bars represent \u00b11 standard error. Asterisks (*) denote significant "
     "between-group contrasts at that session; double asterisks (**) denote ",
     size=10)
_run(p, "p", italic=True, size=10)
_run(p,
     " < .01. CEN = Central Executive Network; DMN = Default Mode Network; "
     "SN = Salience Network.",
     size=10)


# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(str(OUTPUT))
print(f"Saved: {OUTPUT}")
